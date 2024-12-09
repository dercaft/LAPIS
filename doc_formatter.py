from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
from typing import List, Dict

class DocFormatter:
    def __init__(self):
        self.doc = Document()
        self._setup_document()
        
    def _setup_document(self):
        """设置文档默认格式"""
        # 设置默认字体
        self.doc.styles['Normal'].font.name = '宋体'
        self.doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 设置默认段落格式
        self.doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        self.doc.styles['Normal'].paragraph_format.space_before = Pt(6)
        self.doc.styles['Normal'].paragraph_format.space_after = Pt(6)
    
    def add_title(self, title: str):
        """添加文档标题（黑体三号加粗）"""
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run(title)
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(16)  # 三号字体
        run.font.bold = True
        
    def add_chapter(self, title: str):
        """添加章标题（黑体三号）"""
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(title)
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(16)  # 三号字体
        
    def add_section(self, title: str):
        """添加节标题（黑体四号）"""
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(title)
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(14)  # 四号字体
        
    def add_subsection(self, title: str):
        """添加小节标题（宋体小四加粗）"""
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(title)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)  # 小四号字体
        run.font.bold = True
        
    def add_content(self, content: str):
        """添加正文内容（宋体小四）"""
        for line in content.splitlines():  # 按行分割内容
            # 跳过以多个 # 开头的行，跳过空行
            if line.strip().startswith('##') or not line.strip():
                continue
            paragraph = self.doc.add_paragraph()  # 为每一行添加一个新段落
            run = paragraph.add_run(line)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(12)  # 小四号字体
        
    def save(self, filename: str):
        """保存文档"""
        self.doc.save(filename) 